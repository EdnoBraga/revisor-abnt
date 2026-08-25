#!/usr/bin/env python3
"""Create a conservative ABNT-oriented audit report for an existing DOCX."""

import argparse
import json
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def cm(value):
    return round(value.cm, 2) if value is not None else None


def key(text):
    return "".join(c for c in unicodedata.normalize("NFD", text.upper()) if unicodedata.category(c) != "Mn")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()
    source = args.input.resolve()
    out = args.out or source.with_name(source.stem + "_abnt_report.json")
    doc = Document(source)
    headings, quotation_without_locator, citation_hits = [], [], []
    references_index = None
    element_positions = {"sumario": [], "resumo": [], "abstract": [], "keywords": []}
    section_numbers = []
    list_numbered_headings = []
    citation_re = re.compile(r"\((?:[A-Z][A-Z .-]+),\s*(?:19|20)\d{2}(?:[a-z])?(?:,\s*p\.\s*\d+)?\)")
    quote_re = re.compile(r'[\"\u201c\u201d].{15,}[\"\u201c\u201d]')
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""
        if key(text) in {"REFERENCIAS", "REFERENCIAS BIBLIOGRAFICAS", "REFERENCES"}:
            references_index = index
        upper = key(text)
        if upper == "SUMARIO":
            element_positions["sumario"].append(index)
        if upper == "RESUMO":
            element_positions["resumo"].append(index)
        if upper in {"ABSTRACT", "RESUMEN", "RESUME"}:
            element_positions["abstract"].append(index)
        if upper.startswith(("PALAVRAS-CHAVE", "KEYWORDS", "PALABRAS CLAVE", "MOTS-CLES")):
            element_positions["keywords"].append(index)
        if style.startswith("Heading") or re.match(r"^\d+(?:\.\d+)*\s+", text):
            headings.append({"paragraph": index, "text": text[:160], "style": style})
        ppr = paragraph._p.pPr
        has_num_pr = ppr is not None and ppr.find(qn("w:numPr")) is not None
        looks_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
        if has_num_pr and text and len(text) < 100 and (text.isupper() or looks_bold) and "List" in style:
            list_numbered_headings.append({"paragraph": index, "text": text[:160], "style": style, "note": "Uses Word auto-list numbering (numPr), not literal section-number text or a Heading style. The visible number/tab spacing is controlled by the list definition; verify manually that it renders as 'N TITULO' per NBR 6024, since automated formatting cannot safely rewrite list-numbering tab stops."})
        number_match = re.match(r"^(\d+(?:\.\d+)*)\s+", text)
        if number_match:
            section_numbers.append({"paragraph": index, "number": number_match.group(1), "text": text[:160]})
        if citation_re.search(text):
            citation_hits.append(index)
        if quote_re.search(text) and not re.search(r"(?:p\.\s*\d+|p\s+\d+|pagina\s+\d+)", text, re.I):
            quotation_without_locator.append(index)
    sections = []
    for section in doc.sections:
        header_xml = section.header._element.xml
        first_header_xml = section.first_page_header._element.xml
        sections.append({"page_width_cm": cm(section.page_width), "page_height_cm": cm(section.page_height), "top_cm": cm(section.top_margin), "bottom_cm": cm(section.bottom_margin), "left_cm": cm(section.left_margin), "right_cm": cm(section.right_margin), "default_header_has_page_field": "instrText" in header_xml and "PAGE" in header_xml, "first_page_header_has_page_field": "instrText" in first_header_xml and "PAGE" in first_header_xml, "different_first_page_header_footer": section.different_first_page_header_footer})
    manual_validation = ["Confirm institutional manual/template precedence.", "Confirm every direct quote against its original source and locator.", "Confirm reference metadata and citation-reference correspondence.", "Confirm resumo/abstract content, keywords, and length rule.", "Update and verify sumario fields in Word after formatting.", "Review tables, figures, front matter, and page numbering visually."]
    if list_numbered_headings:
        manual_validation.append("Section titles appear to use Word's auto-list numbering (List Paragraph/numPr) instead of a Heading style or literal numbers; see list_numbered_heading_candidates. The formatter does not rewrite list-numbering tab stops, so confirm visually that the number-to-title spacing matches NBR 6024.")
    report = {"input": str(source), "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables), "section_count": len(doc.sections), "sections": sections, "references_heading_paragraph": references_index, "structural_element_positions": element_positions, "progressive_section_number_candidates": section_numbers, "heading_candidates": headings, "list_numbered_heading_candidates": list_numbered_headings, "author_date_citation_candidates": citation_hits, "quoted_paragraphs_without_locator_candidates": quotation_without_locator, "manual_validation": manual_validation}
    out.write_text(json.dumps(report, ensure_ascii=True, indent=2), encoding="utf-8")
    print(out)


if __name__ == "__main__":
    main()
