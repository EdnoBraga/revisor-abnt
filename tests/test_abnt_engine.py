"""Teste de regressao do motor: garante que a formatacao e de fato aplicada.

Este teste existe porque o app ja ficou, em producao, rodando uma imagem
Docker desatualizada em relacao ao motor commitado -- o usuario via "revisao
concluida" mas o arquivo voltava praticamente identico ao original. Um teste
que falha sempre que o motor para de mudar um documento sintetico conhecido
detecta esse tipo de regressao antes da implantacao, nao depois.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "revisor-abnt-docx" / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from abnt_engine import ReviewConfig, add_word_field, apply_formatting, page_field_count, reference_title_span, scan_document  # noqa: E402


def _add(doc, text, style=None, alignment=None, indent_cm=None):
    paragraph = doc.add_paragraph(text, style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    if indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(indent_cm)
    return paragraph


def _build_tcc_docx(path: Path, *, explicit_introducao: bool) -> None:
    """Documento sintetico com um paragrafo de corpo deliberadamente mal
    formatado (centralizado, sem recuo), para que o teste consiga detectar
    se o motor realmente o corrige."""
    doc = Document()
    _add(doc, "NOME DO ALUNO", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, "TITULO DO TRABALHO", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "Palavras-chave: um; dois; tres.")

    heading_text = "1 INTRODUÇÃO" if explicit_introducao else "1 USO DA INTELIGÊNCIA ARTIFICIAL"
    _add(doc, heading_text, style="Heading 1")
    # Paragrafo de corpo fora do padrao ABNT: centralizado, sem recuo de 1a linha.
    body = _add(
        doc,
        "Este e um paragrafo de corpo que deveria ficar justificado, com recuo de "
        "primeira linha de 1,25 cm, mas comeca centralizado e sem recuo (SILVA, 2020).",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=0,
    )
    assert body.alignment == WD_ALIGN_PARAGRAPH.CENTER

    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")
    doc.save(path)


def test_apply_formatting_changes_a_misformatted_body_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    _build_tcc_docx(source, explicit_introducao=True)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc"))

    # A rede de seguranca central deste teste: se o motor parar de agir,
    # actions_applied fica vazio e este assert falha.
    codes = {action["code"] for action in report["actions_applied"]}
    assert "formatted_body" in codes, report["actions_applied"]

    revised = Document(output)
    body_paragraph = next(p for p in revised.paragraphs if p.text.startswith("Este e um paragrafo de corpo"))
    assert body_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    # Cm(1.25) e o first_line_indent lido de volta podem diferir em fracoes de
    # EMU por arredondamento interno do docx; comparar em cm evita esse ruido.
    assert round(body_paragraph.paragraph_format.first_line_indent.cm, 2) == 1.25


def test_textual_start_falls_back_to_first_heading_without_toc_or_literal_title(tmp_path: Path) -> None:
    """Documentos sem SUMARIO e com capitulos nomeados pelo tema (nao
    literalmente 'INTRODUCAO') ainda devem ter o corpo formatado."""
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    _build_tcc_docx(source, explicit_introducao=False)

    scan = scan_document(Document(source), ReviewConfig(document_type="tcc"))
    assert scan["textual_start_detection"] == "heading_fallback"
    assert scan["textual_start_paragraph"] is not None

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "formatted_body" in codes, report["actions_applied"]


def test_institution_profile_is_recorded_in_the_report(tmp_path: Path) -> None:
    source = tmp_path / "original.docx"
    _build_tcc_docx(source, explicit_introducao=True)
    scan = scan_document(Document(source), ReviewConfig(document_type="article", institution="cgaem"))
    assert scan["institution"] == "cgaem"


def test_reference_title_span_bolds_book_title_only() -> None:
    text = "SILVA, João. Manual de metodologia científica. São Paulo: Atlas, 2020."
    span = reference_title_span(text)
    assert span is not None
    assert text[span[0]:span[1]] == "Manual de metodologia científica."


def test_reference_title_span_bolds_nbr_standard_number() -> None:
    text = (
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 6023: informação e "
        "documentação: referências: elaboração. Rio de Janeiro: ABNT, 2018."
    )
    span = reference_title_span(text)
    assert span is not None
    assert text[span[0]:span[1]] == "NBR 6023"


def test_reference_title_span_bolds_legal_instrument_name() -> None:
    text = (
        "BRASIL. Lei nº 9.394, de 20 de dezembro de 1996. Estabelece as diretrizes "
        "e bases da educação nacional. Brasília, DF: Presidência da República, 1996."
    )
    span = reference_title_span(text)
    assert span is not None
    assert text[span[0]:span[1]] == "Lei nº 9.394, de 20 de dezembro de 1996"


def test_reference_title_span_skips_periodical_article() -> None:
    text = (
        "CÁCERES, Maria. Título de artigo qualquer. Revista Brasileira de "
        "Educação, v. 10, n. 2, p. 1-20, 2019."
    )
    assert reference_title_span(text) is None


def test_reference_title_span_skips_book_chapter_with_in() -> None:
    text = (
        "FINK, Daniel. Título do capítulo. In: SOUZA, Ana (org.). Título do "
        "livro. São Paulo: Vozes, 2015. p. 10-20."
    )
    assert reference_title_span(text) is None


def test_apply_formatting_bolds_reference_titles_end_to_end(tmp_path: Path) -> None:
    """Reproduz a queixa do usuario: 'a referencia burocratica tinha a parte que
    nao estava em negrito'. O motor precisa de fato negritar o titulo, nao so
    normalizar espacamento/alinhamento."""
    doc = Document()
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Manual de metodologia científica. São Paulo: Atlas, 2020.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "reference_titles_bolded" in codes, report["actions_applied"]

    revised = Document(output)
    reference_paragraph = next(p for p in revised.paragraphs if p.text.startswith("SILVA, João"))
    bold_text = "".join(run.text for run in reference_paragraph.runs if run.bold)
    assert bold_text == "Manual de metodologia científica."


def test_apply_formatting_creates_sumario_when_missing_entirely(tmp_path: Path) -> None:
    """Reproduz a outra queixa do usuario: 'ele nao acertou a parte do sumario'.
    Quando o documento nao tem nenhum SUMARIO e o modo insert-if-empty esta
    ativo, o motor precisa criar a secao com um campo TOC nativo, nao so
    deixar passar em branco."""
    doc = Document()
    _add(doc, "NOME DO ALUNO", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc", toc_mode="insert-if-empty"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "sumario_section_inserted" in codes, report["actions_applied"]

    revised = Document(output)
    texts = [p.text for p in revised.paragraphs]
    sumario_index = texts.index("SUMÁRIO")
    introducao_index = next(i for i, t in enumerate(texts) if t.startswith("1 INTRODUÇÃO"))
    assert sumario_index < introducao_index
    # O campo TOC nativo fica no parágrafo logo após o título SUMÁRIO.
    field_paragraph = revised.paragraphs[sumario_index + 1]
    assert "TOC" in field_paragraph._p.xml.upper()


def test_apply_formatting_does_not_touch_toc_in_audit_mode(tmp_path: Path) -> None:
    """Sem toc_mode='insert-if-empty', o motor so audita: nao deve inserir
    nenhuma secao SUMARIO nem alterar o documento por conta propria."""
    doc = Document()
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc", toc_mode="audit"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "sumario_section_inserted" not in codes


def test_apply_formatting_fixes_case_in_apud_indirect_citations(tmp_path: Path) -> None:
    """Reproduz a terceira queixa do usuario: citacoes indiretas (apud) nao
    eram corrigidas -- nem o autor original nem o autor citante tinham a
    caixa alta convertida, porque a palavra "apud" quebrava as duas regras
    de deteccao existentes (bloco de autoria deixa de ser 100% maiusculo, ou
    o autor original nao fica colado no "(ano)")."""
    doc = Document()
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Segundo SILVA (1990 apud SOUZA, 2020), o tema é relevante para a área.")
    _add(doc, "Já PIAGET apud VYGOTSKY (1978) defendia outra tese.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SOUZA, Ana. Título do trabalho. Cidade: Editora, 2020.")
    _add(doc, "VYGOTSKY, Lev. Outro título. Cidade: Editora, 1978.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc", citation_system="author-date"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "citation_author_case_normalized" in codes, report["actions_applied"]

    revised = Document(output)
    texts = "\n".join(p.text for p in revised.paragraphs)
    assert "Silva (1990 apud Souza, 2020)" in texts
    assert "Piaget apud Vygotsky (1978)" in texts


def test_apply_formatting_applies_pagination_section_break_by_default(tmp_path: Path) -> None:
    """Reproduz a queixa do usuario sobre paginacao: 'e pra numerar
    corretamente, conforme ABNT'. Com pagination_mode='request' (o padrao
    desde que o usuario pediu correcao em vez de auditoria), o motor precisa
    de fato inserir a quebra de secao antes do inicio textual e um campo
    PAGE nativo no cabecalho da nova secao -- sem numerar capa/pre-textuais
    nem reiniciar a contagem."""
    doc = Document()
    _add(doc, "NOME DO ALUNO", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc"))
    assert report["config"]["pagination_mode"] == "request"
    codes = {action["code"] for action in report["actions_applied"]}
    assert "pagination_section_break_inserted" in codes, report["actions_applied"]

    revised = Document(output)
    assert len(revised.sections) == 2
    section_pretextual, section_textual = revised.sections
    assert page_field_count(section_pretextual) == 0
    assert section_textual.start_type == WD_SECTION_START.NEW_PAGE
    assert page_field_count(section_textual) > 0


def test_apply_formatting_replaces_pre_existing_page_field(tmp_path: Path) -> None:
    """Reproduz a queixa do usuario: 'quero que atualize, independente de como
    estiver o arquivo de origem, pois o motivo da aplicacao e de fato corrigir'.
    Um documento que ja chega com um campo PAGE no cabecalho global (comum em
    modelos institucionais que numeram tudo, inclusive capa/pre-textuais) nao
    pode mais ser deixado de lado -- o motor precisa substituir esse cabecalho
    pela numeracao correta, comecando so na parte textual."""
    doc = Document()
    _add(doc, "NOME DO ALUNO", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")

    # Simula um modelo institucional que ja numera a partir da capa (errado pela
    # NBR 14724, mas um cenario real e comum).
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    existing_header_paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    existing_header_paragraph.add_run("Página ")
    add_word_field(existing_header_paragraph, "PAGE")

    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)
    assert page_field_count(Document(source).sections[0]) > 0  # pré-condição do teste

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "pagination_section_break_inserted" in codes, report["actions_applied"]
    pagination_action = next(a for a in report["actions_applied"] if a["code"] == "pagination_section_break_inserted")
    assert "substituíd" in pagination_action["message"]

    revised = Document(output)
    assert len(revised.sections) == 2
    section_pretextual, section_textual = revised.sections
    # A seção pré-textual (capa, resumo) não pode mais mostrar número algum,
    # mesmo tendo herdado o cabeçalho antigo que numerava tudo.
    assert page_field_count(section_pretextual) == 0
    # A seção textual tem exatamente um campo PAGE (o antigo foi removido, não
    # apenas complementado por um novo).
    header_xml = section_textual.header._element.xml.upper()
    assert header_xml.count("PAGE") == 1
    assert "PÁGINA" not in header_xml  # texto antigo "Página " não sobrou


def test_apply_formatting_does_not_touch_pagination_in_audit_mode(tmp_path: Path) -> None:
    """Sem pagination_mode='request', o motor so audita: o documento deve
    permanecer com uma unica secao, sem quebra nem campo PAGE inseridos."""
    doc = Document()
    _add(doc, "RESUMO", style="Heading 1")
    _add(doc, "Resumo do trabalho com o numero minimo de palavras necessario " * 10)
    _add(doc, "1 INTRODUÇÃO", style="Heading 1")
    _add(doc, "Texto de corpo qualquer para preencher a introdução do documento.")
    _add(doc, "REFERÊNCIAS", style="Heading 1")
    _add(doc, "SILVA, João. Título do trabalho. Cidade: Editora, 2020.")
    source = tmp_path / "original.docx"
    output = tmp_path / "revisado.docx"
    doc.save(source)

    report = apply_formatting(source, output, ReviewConfig(document_type="tcc", pagination_mode="audit"))
    codes = {action["code"] for action in report["actions_applied"]}
    assert "pagination_section_break_inserted" not in codes
    revised = Document(output)
    assert len(revised.sections) == 1


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-v"]))
