import json
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
import sys
import uuid

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi.testclient import TestClient

from app.main import app


SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "revisor-abnt-docx" / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))
from abnt_engine import ReviewConfig, apply_formatting, scan_document  # noqa: E402


def sample_docx() -> bytes:
    document = Document()
    document.add_paragraph("CAPA")
    document.add_heading("1 INTRODUÇÃO", level=1)
    document.add_paragraph("Texto de teste para revisão.")
    document.add_heading("REFERÊNCIAS", level=1)
    document.add_paragraph("SILVA, Maria. Título de teste. Brasília: Editora, 2026.")
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def test_health():
    client = TestClient(app)
    assert client.get("/api/health").json()["status"] == "ok"


def test_revision_requires_privacy_acknowledgement():
    client = TestClient(app)
    response = client.post(
        "/api/revisions",
        files={"document": ("tcc.docx", sample_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 422


def test_docx_revision_creates_downloadable_copy(tmp_path, monkeypatch):
    monkeypatch.setenv("REVISOR_ABNT_JOBS_DIR", str(tmp_path / "jobs"))
    client = TestClient(app)
    response = client.post(
        "/api/revisions",
        files={"document": ("tcc.docx", sample_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={
            "privacy_acknowledged": "true",
            "document_type": "tcc",
            "citation_system": "author-date",
            "font": "Arial",
            "toc_mode": "audit",
            "order_references": "true",
        },
    )
    assert response.status_code == 202
    job = client.get(response.json()["status_url"]).json()
    assert job["status"] == "completed"
    download = client.get(job["download_url"])
    assert download.status_code == 200
    assert download.content[:2] == b"PK"
    report = client.get(job["format_report_url"])
    assert report.status_code == 200
    assert report.json()["config"]["font"] == "Arial"
    assert job["review_summary"]["actions_applied"] > 0

    # O usuario pediu explicitamente um relatorio em PDF, nao apenas JSON, com
    # o que foi corrigido automaticamente -- este teste garante que o endpoint
    # existe e entrega um PDF de verdade, nao apenas um link quebrado.
    pdf_report = client.get(job["format_report_pdf_url"])
    assert pdf_report.status_code == 200
    assert pdf_report.headers["content-type"] == "application/pdf"
    assert pdf_report.content[:4] == b"%PDF"


def test_default_options_apply_corrections_instead_of_only_auditing(tmp_path, monkeypatch):
    """Reproduz o pedido do usuario: 'eu nao quero auditagem, eu quero
    correcao'. Sem o usuario escolher nada nas opcoes de revisao, sumario e
    paginacao devem ser corrigidos automaticamente, nao apenas auditados."""
    monkeypatch.setenv("REVISOR_ABNT_JOBS_DIR", str(tmp_path / "jobs"))
    client = TestClient(app)
    response = client.post(
        "/api/revisions",
        files={"document": ("tcc.docx", sample_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"privacy_acknowledged": "true"},
    )
    assert response.status_code == 202
    job = client.get(response.json()["status_url"]).json()
    assert job["status"] == "completed"
    report = client.get(job["format_report_url"]).json()
    assert report["config"]["toc_mode"] == "insert-if-empty"
    assert report["config"]["pagination_mode"] == "request"
    codes = {action["code"] for action in report["actions_applied"]}
    assert "sumario_section_inserted" in codes, report["actions_applied"]
    assert "pagination_section_break_inserted" in codes, report["actions_applied"]


def test_engine_applies_real_layout_and_preserves_word_numbering(tmp_path):
    source = tmp_path / "entrada.docx"
    output = tmp_path / "saida.docx"
    document = Document()
    document.add_paragraph("CAPA - NÃO FORMATAR COMO CORPO")
    document.add_paragraph("SUMÁRIO")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")
    heading = document.add_heading("1 INTRODUÇÃO", level=1)
    numbering = OxmlElement("w:numPr")
    heading._p.get_or_add_pPr().append(numbering)
    body = document.add_paragraph("Texto normal que deve receber formatação de corpo.")
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.paragraph_format.first_line_indent = Cm(0)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    body.runs[0].font.name = "Arial"
    body.runs[0].font.size = Pt(10)
    quote = document.add_paragraph("Este é um bloco longo identificável pela apresentação já existente (Silva, 2020, p. 12).")
    quote.paragraph_format.left_indent = Cm(4)
    document.add_paragraph("“Trecho direto sem indicação de página para auditoria.”")
    document.add_paragraph("REFERÊNCIAS")
    document.add_paragraph("ZETA, Zoe. Obra Z. Cidade: Editora, 2020.")
    document.add_paragraph("ALFA, Ana. Obra A. Cidade: Editora, 2019.")
    document.save(source)

    report = apply_formatting(source, output, ReviewConfig(toc_mode="insert-if-empty", citation_system="author-date"))
    revised = Document(output)
    body_out = next(p for p in revised.paragraphs if p.text.startswith("Texto normal"))
    quote_out = next(p for p in revised.paragraphs if p.text.startswith("Este é um bloco"))
    heading_out = next(p for p in revised.paragraphs if p.text == "1 INTRODUÇÃO")
    references_heading = next(p for p in revised.paragraphs if p.text == "REFERÊNCIAS")
    references = [p.text for p in revised.paragraphs if p.text.startswith(("ALFA,", "ZETA,"))]

    assert round(revised.sections[0].top_margin.cm, 2) == 3
    assert round(revised.sections[0].left_margin.cm, 2) == 3
    assert round(revised.sections[0].right_margin.cm, 2) == 2
    assert body_out.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert round(body_out.paragraph_format.first_line_indent.cm, 2) == 1.25
    assert body_out.runs[0].font.name == "Times New Roman"
    assert body_out.runs[0].font.size.pt == 12
    assert round(quote_out.paragraph_format.left_indent.cm, 2) == 4
    assert quote_out.runs[0].font.size.pt == 10
    assert heading_out._p.pPr.find(qn("w:numPr")) is not None
    assert references_heading.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert references == ["ALFA, Ana. Obra A. Cidade: Editora, 2019.", "ZETA, Zoe. Obra Z. Cidade: Editora, 2020."]
    assert "TOC" in revised.element.body.xml
    assert any(action["code"] == "references_reordered_author_date" for action in report["actions_applied"])
    assert any(issue["code"] == "direct_quote_without_locator" for issue in report["issues_remaining"])


def test_engine_normalizes_author_case_only_in_textual_citations(tmp_path):
    source = tmp_path / "entrada-citacoes.docx"
    output = tmp_path / "saida-citacoes.docx"
    document = Document()
    document.add_paragraph("CAPA")
    document.add_heading("1 INTRODUÇÃO", level=1)
    document.add_paragraph("A análise considera (SILVA; SOUZA, 2020) e (IBGE, 2021). SILVA (2020) também discute o tema.")
    document.add_heading("REFERÊNCIAS", level=1)
    document.add_paragraph("SILVA, Maria. Obra de Silva. Brasília: Editora, 2020.")
    document.add_paragraph("SOUZA, João. Obra de Souza. Brasília: Editora, 2020.")
    document.add_paragraph("IBGE. Pesquisa de referência. Rio de Janeiro: IBGE, 2021.")
    document.save(source)

    report = apply_formatting(source, output, ReviewConfig(citation_system="author-date"))
    revised = Document(output)
    body = next(p.text for p in revised.paragraphs if p.text.startswith("A análise considera"))
    references = [p.text for p in revised.paragraphs if p.text.startswith(("SILVA,", "SOUZA,", "IBGE."))]

    assert "(Silva; Souza, 2020)" in body
    assert "(IBGE, 2021)" in body
    assert "Silva (2020)" in body
    assert set(references) == {
        "SILVA, Maria. Obra de Silva. Brasília: Editora, 2020.",
        "SOUZA, João. Obra de Souza. Brasília: Editora, 2020.",
        "IBGE. Pesquisa de referência. Rio de Janeiro: IBGE, 2021.",
    }
    assert any(action["code"] == "citation_author_case_normalized" for action in report["actions_applied"])


def test_engine_preserves_citation_split_between_runs(tmp_path):
    source = tmp_path / "entrada-citacao-dividida.docx"
    output = tmp_path / "saida-citacao-dividida.docx"
    document = Document()
    document.add_heading("1 INTRODUÇÃO", level=1)
    paragraph = document.add_paragraph("Texto com ")
    paragraph.add_run("(SIL").bold = True
    paragraph.add_run("VA, 2020) em trechos distintos.")
    document.add_heading("REFERÊNCIAS", level=1)
    document.add_paragraph("SILVA, Maria. Obra de Silva. Brasília: Editora, 2020.")
    document.save(source)

    report = apply_formatting(source, output, ReviewConfig(citation_system="author-date"))
    revised = Document(output)
    body = next(p for p in revised.paragraphs if p.text.startswith("Texto com"))

    assert body.text == "Texto com (SILVA, 2020) em trechos distintos."
    assert body.runs[1].bold is True
    assert any(issue["code"] == "citation_case_split_across_runs" for issue in report["issues_remaining"])


def test_cleanup_removes_stale_processing_jobs(tmp_path, monkeypatch):
    """Um job preso em 'processing' (processo morto por crash/reinicio) precisa
    ser reaproveitado pela limpeza, senao o diretorio de jobs cresce para sempre
    -- BackgroundTasks nao tem supervisor nem retomada propria."""
    monkeypatch.setenv("REVISOR_ABNT_JOBS_DIR", str(tmp_path / "jobs"))
    from app.services import document_processor as dp

    stale_id = str(uuid.uuid4())
    stale_dir = tmp_path / "jobs" / stale_id
    stale_dir.mkdir(parents=True)
    stale_created = (datetime.now(UTC) - timedelta(minutes=45)).isoformat()
    (stale_dir / "job.json").write_text(
        json.dumps({"id": stale_id, "status": "processing", "created_at": stale_created}),
        encoding="utf-8",
    )

    fresh_id = str(uuid.uuid4())
    fresh_dir = tmp_path / "jobs" / fresh_id
    fresh_dir.mkdir(parents=True)
    fresh_created = datetime.now(UTC).isoformat()
    (fresh_dir / "job.json").write_text(
        json.dumps({"id": fresh_id, "status": "processing", "created_at": fresh_created}),
        encoding="utf-8",
    )

    removed = dp.cleanup_expired_jobs()

    assert removed == 1
    assert not stale_dir.exists()
    assert fresh_dir.exists()


def test_audit_does_not_claim_reference_or_pagination_corrections_without_evidence():
    document = Document()
    document.add_heading("1 INTRODUÇÃO", level=1)
    document.add_paragraph("Citação possível (Silva, 2020).")
    report = scan_document(document, ReviewConfig())
    issue_codes = {issue["code"] for issue in report["issues"]}

    assert "citations_without_reference_list" in issue_codes
    assert "page_number_field_not_found" in issue_codes
    assert report["citation_system"] == "author-date"
