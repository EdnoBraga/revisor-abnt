"""Motor conservador de auditoria e formatação ABNT para arquivos DOCX.

O motor aplica somente regras que podem ser verificadas na estrutura do Word.
Ele não completa metadados, não cria citações e não altera o sentido de textos.
"""

from __future__ import annotations

import copy
import re
import unicodedata
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


DEFAULT_FONT = "Times New Roman"
TEXTUAL_STARTS = {"INTRODUCAO", "APRESENTACAO", "DESENVOLVIMENTO"}
REFERENCE_HEADINGS = {"REFERENCIAS", "REFERENCIAS BIBLIOGRAFICAS", "REFERENCES"}
POST_TEXTUAL_STARTS = ("APENDICE", "ANEXO", "GLOSSARIO", "INDICE")
ABSTRACT_HEADINGS = {"RESUMO", "ABSTRACT", "RESUMEN", "RESUME"}
KEYWORD_PREFIXES = ("PALAVRAS-CHAVE", "PALAVRAS CHAVE", "KEYWORDS", "PALABRAS CLAVE", "MOTS-CLES")
UNNUMBERED_HEADINGS = {
    "AGRADECIMENTOS", "RESUMO", "ABSTRACT", "RESUMEN", "RESUME", "LISTA DE ILUSTRACOES",
    "LISTA DE TABELAS", "LISTA DE ABREVIATURAS E SIGLAS", "LISTA DE SIMBOLOS", "SUMARIO",
    "REFERENCIAS", "REFERENCIAS BIBLIOGRAFICAS", "GLOSSARIO", "INDICE",
}
SECTION_RE = re.compile(r"^\s*(\d+(?:\.\d+){0,4})\.?\s+\S")
YEAR_RE = re.compile(r"\b((?:1[5-9]|20)\d{2}[a-z]?)\b", re.IGNORECASE)
PARENTHETICAL_RE = re.compile(r"\(([^()]{1,220})\)")
QUOTED_RE = re.compile(r"[\"\u201c\u201d].{15,}[\"\u201c\u201d]", re.DOTALL)
LOCATOR_RE = re.compile(r"\b(?:p\.|pagina|página|localiz(?:acao|ação)|posi(?:cao|ção))\s*\d+", re.IGNORECASE)
PARENTHETICAL_AUTHOR_YEAR_RE = re.compile(
    r"(?:^|(?P<separator>;))(?P<leading>\s*)(?P<authors>[^,()]{2,100}?)\s*,\s*(?P<year>(?:1[5-9]|20)\d{2}[a-z]?)\b",
    re.IGNORECASE,
)
NARRATIVE_AUTHOR_YEAR_RE = re.compile(
    r"\b(?P<authors>[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ý'’.-]*(?:\s+(?:[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ý'’.-]*|E|DA|DAS|DE|DO|DOS))*?(?:\s+ET\s+AL\.)?)\s*\(\s*(?P<year>(?:1[5-9]|20)\d{2}[a-z]?)\b"
)
# Citação indireta (NBR 10520): "SOBRENOME-ORIGINAL apud SOBRENOME-CITANTE, ano".
# O nome logo depois de "apud" segue o mesmo padrão autor-vírgula-ano das
# citações comuns; o nome logo antes de "apud" não tem ano adjacente (o ano
# citado é sempre o da obra consultada diretamente), por isso é tratado à
# parte, exigindo apenas que seja uma sequência de palavras em caixa alta.
APUD_CITING_AUTHOR_RE = re.compile(
    r"(?P<apud>(?i:apud))(?P<sep>\s+)(?P<authors>[^,()]{2,100}?)\s*,\s*(?P<year>(?:1[5-9]|20)\d{2}[a-z]?)\b"
)
APUD_ORIGINAL_AUTHOR_RE = re.compile(
    r"\b(?P<authors>[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ý'’.-]*(?:\s+(?:[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ý'’.-]*|E|DA|DAS|DE|DO|DOS))*?)(?P<sep>\s+)(?P<apud>(?i:apud))\b"
)
NAME_PARTICLES = {"da", "das", "de", "do", "dos", "e"}
# Siglas são uma exceção à forma "Maiúscula/minúscula" da NBR 10520:2023.
# A lista reduz falsos positivos quando a referência ainda não permite identificar
# com segurança se uma chamada curta é sobrenome ou entidade siglada.
KNOWN_INITIALISMS = {
    "ABNT", "ANEEL", "ANVISA", "CAPES", "CGU", "CNPQ", "IBAMA", "IBGE", "IPEA",
    "MEC", "OCDE", "OIT", "OMS", "ONU", "STF", "STJ", "TCU", "UNESCO", "UNICEF",
}
# Destaque tipográfico do título em referências (NBR 6023): não há uma regra fixa
# universal, mas o manual institucional do CGAEM/ESFCEx usado como referência
# aplica negrito de forma consistente ao título de livros/monografias, ao número
# de normas NBR e ao nome de instrumentos legais. As expressões abaixo só
# reconhecem esses casos concretos e devolvem None (sem alterar nada) sempre que
# o formato da entrada foge do que foi validado nos exemplos do manual —
# preferindo deixar uma entrada sem negrito automático a arriscar negritar o
# trecho errado.
AUTHOR_BLOCK_RE = re.compile(
    r"^\s*(?:[A-ZÀ-Ý][A-Za-zÀ-ÿ'’\-]*(?:\s+[A-Za-zÀ-ÿ'’\-]+)*\s*,\s*[^;.]+?(?:;\s*)?)+\.\s+"
)
LEGAL_INSTRUMENT_RE = re.compile(
    r"(Decreto(?:-Lei)?|Lei\s+Complementar|Lei|Portaria|Resolu[çc][ãa]o|Instru[çc][ãa]o\s+Normativa|"
    r"Medida\s+Provis[óo]ria|Emenda\s+Constitucional|Decreto\s+Legislativo)\s+n[ºo°]?\s*[\d.\/-]+"
    r"(?:,\s*de\s+\d{1,2}\s+de\s+[a-zà-ú]+\s+de\s+\d{4})?",
    re.IGNORECASE,
)
NBR_STANDARD_RE = re.compile(r"ASSOCIA[ÇC][ÃA]O BRASILEIRA DE NORMAS T[ÉE]CNICAS\.\s+(NBR\s*[\d.\-]+)", re.IGNORECASE)
PERIODICAL_MARKER_RE = re.compile(r",\s*v\.\s*\d+|,\s*n\.\s*\d+")
TAIL_PLACE_PUBLISHER_YEAR_RE = re.compile(r":\s*([^,:]{2,100}),\s*(?:19|20)\d{2}[a-z]?\.?\s*$")
TITLE_BOUNDARY_RE = re.compile(r"[.?!](?=\s)")


@dataclass(frozen=True)
class ReviewConfig:
    document_type: str = "tcc"
    citation_system: str = "auto"
    font: str = DEFAULT_FONT
    toc_mode: str = "audit"
    pagination_mode: str = "audit"
    order_references: bool = True
    institution: str = "generic"


def normalized(text: str) -> str:
    """Normaliza títulos para comparação, sem usar essa forma para reescrever texto."""
    decomposed = unicodedata.normalize("NFD", text.upper().strip())
    return " ".join("".join(c for c in decomposed if unicodedata.category(c) != "Mn").split())


def cm(value) -> float | None:
    return round(value.cm, 2) if value is not None else None


def paragraph_numbering(paragraph: Paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    style_match = re.search(r"(?:Heading|Titulo|Título)\s*([1-5])$", style_name, re.IGNORECASE)
    if style_match:
        return int(style_match.group(1))
    text_match = SECTION_RE.match(paragraph.text)
    if text_match:
        return text_match.group(1).count(".") + 1
    return None


def is_semantic_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    return heading_level(paragraph) is not None or normalized(text) in UNNUMBERED_HEADINGS


def is_reference_heading(text: str) -> bool:
    return normalized(text) in REFERENCE_HEADINGS


def is_post_textual_heading(text: str) -> bool:
    return normalized(text).startswith(POST_TEXTUAL_STARTS)


def set_font(run, name: str, size: int) -> None:
    """Ajusta fonte/tamanho mantendo negrito, itálico, sublinhado e cor existentes."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def set_style_font(style, name: str, size: int) -> None:
    style.font.name = name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)


def set_paragraph_font(paragraph: Paragraph, name: str, size: int) -> None:
    for run in paragraph.runs:
        if run.text:
            set_font(run, name, size)


def add_issue(issues: list[dict[str, Any]], *, code: str, message: str, severity: str = "warning", paragraph: int | None = None, auto_fixable: bool = False) -> None:
    issues.append(
        {
            "code": code,
            "severity": severity,
            "paragraph": paragraph,
            "message": message,
            "auto_fixable": auto_fixable,
        }
    )


def _citation_parts(text: str) -> list[dict[str, str]]:
    """Extrai somente candidatos autor-data; não decide se uma citação é válida."""
    parts: list[dict[str, str]] = []
    for match in PARENTHETICAL_RE.finditer(text):
        content = match.group(1)
        if not YEAR_RE.search(content):
            continue
        # Um bloco pode conter várias obras separadas por ponto e vírgula.
        for item in content.split(";"):
            candidate = re.search(r"^\s*([^,()]{2,100}?)(?:,\s*)((?:1[5-9]|20)\d{2}[a-z]?)\b", item, re.IGNORECASE)
            if not candidate:
                continue
            author = candidate.group(1).strip()
            if author.lower().startswith(("p.", "pagina", "página")):
                continue
            parts.append(
                {
                    "author": author,
                    "year": candidate.group(2).lower(),
                    "text": match.group(0),
                }
            )
    return parts


def citation_key(author: str, year: str) -> str:
    # No sistema autor-data a entrada normalmente é recuperável pelo primeiro sobrenome ou entidade.
    clean = normalized(author).replace("ET AL.", "").strip()
    clean = re.split(r"\s+(?:E|AND)\s+|;", clean)[0].strip()
    return f"{clean}|{year.lower()}"


def reference_key(text: str) -> str | None:
    clean = text.strip()
    year = YEAR_RE.search(clean)
    if not year:
        return None
    # Para entradas ABNT tradicionais, a primeira parte antes de vírgula/ponto é a entrada de autoria.
    author_match = re.match(r"^\s*([^,\.]{2,100})(?:,|\.)", clean)
    if not author_match:
        return None
    author = normalized(author_match.group(1))
    if not author:
        return None
    return f"{author}|{year.group(1).lower()}"


def _reference_author_profiles(paragraphs: list[Paragraph], start: int | None, end: int | None) -> dict[str, set[str]]:
    """Extrai perfis mínimos da lista final sem alterar as referências.

    O sobrenome antes da vírgula costuma indicar pessoa física; entradas curtas
    antes de ponto podem representar entidade siglada. Esses perfis tornam a
    alteração da citação menos arriscada, especialmente para sobrenomes curtos.
    """
    people: set[str] = set()
    initialisms = set(KNOWN_INITIALISMS)
    if start is None or end is None:
        return {"people": people, "initialisms": initialisms}
    for paragraph in paragraphs[start + 1:end]:
        text = paragraph.text.strip()
        match = re.match(r"^\s*([^,\.]{2,100})(?P<separator>,|\.)", text)
        if not match:
            continue
        entry = match.group(1).strip()
        key = normalized(entry)
        if match.group("separator") == ",":
            people.add(key)
        elif re.fullmatch(r"[A-Z]{2,12}", entry):
            initialisms.add(key)
    return {"people": people, "initialisms": initialisms}


def _title_case_name(value: str) -> str:
    """Converte sobrenomes em caixa alta, preservando partículas e hífen/apóstrofo."""
    def title_word(word: str) -> str:
        lower = word.lower()
        if lower in NAME_PARTICLES:
            return lower
        return re.sub(
            r"(^|[-'’])([a-zà-öø-ÿ])",
            lambda match: f"{match.group(1)}{match.group(2).upper()}",
            lower,
        )

    return " ".join(title_word(word) for word in value.split())


def _format_citation_author(author: str, profiles: dict[str, set[str]]) -> tuple[str, bool, bool]:
    """Retorna texto, se houve mudança e se a chamada curta ficou ambígua."""
    match = re.match(r"^(?P<leading>\s*)(?P<core>.*?)(?P<trailing>\s*)$", author, re.DOTALL)
    assert match is not None
    leading, core, trailing = match.group("leading"), match.group("core"), match.group("trailing")
    suffix_match = re.search(r"\s+ET\s+AL\.\s*$", core, re.IGNORECASE)
    suffix = ""
    if suffix_match:
        suffix = " et al."
        core = core[:suffix_match.start()].rstrip()
    if not core or not core.isupper():
        return author, False, False

    key = normalized(core)
    if key in profiles["initialisms"]:
        return author, False, False
    # Uma entrada "SILVA, Nome" na lista final comprova que SILVA é pessoa,
    # mesmo quando o sobrenome é curto. Sem esse indício, siglas curtas ficam
    # para conferência humana; não as transformamos às cegas.
    unspaced = key.replace(" ", "")
    if key not in profiles["people"] and " " not in core and len(unspaced) <= 4:
        return author, False, True

    formatted = _title_case_name(core) + suffix
    return f"{leading}{formatted}{trailing}", formatted != core + suffix, False


def _format_citation_author_expression(authors: str, profiles: dict[str, set[str]]) -> tuple[str, int, int]:
    """Formata uma ou mais autorias separadas por ponto e vírgula."""
    changed = 0
    ambiguous = 0
    chunks = re.split(r"(;)", authors)
    for index in range(0, len(chunks), 2):
        formatted, did_change, is_ambiguous = _format_citation_author(chunks[index], profiles)
        chunks[index] = formatted
        changed += int(did_change)
        ambiguous += int(is_ambiguous)
    return "".join(chunks), changed, ambiguous


def _normalize_citation_case_in_text(text: str, profiles: dict[str, set[str]]) -> tuple[str, int, int]:
    """Normaliza somente chamadas autor-data em um trecho de texto seguro."""
    changes = 0
    ambiguous = 0

    def replace_parenthetical(match: re.Match[str]) -> str:
        nonlocal changes, ambiguous
        content = match.group(1)
        if not _citation_parts(match.group(0)):
            return match.group(0)

        def replace_author_year(author_year_match: re.Match[str]) -> str:
            nonlocal changes, ambiguous
            authors, changed_here, ambiguous_here = _format_citation_author_expression(author_year_match.group("authors"), profiles)
            changes += changed_here
            ambiguous += ambiguous_here
            separator = author_year_match.group("separator") or ""
            return f"{separator}{author_year_match.group('leading')}{authors}, {author_year_match.group('year')}"

        return f"({PARENTHETICAL_AUTHOR_YEAR_RE.sub(replace_author_year, content)})"

    normalized_text = PARENTHETICAL_RE.sub(replace_parenthetical, text)

    def replace_narrative(match: re.Match[str]) -> str:
        nonlocal changes, ambiguous
        authors, changed_here, ambiguous_here = _format_citation_author_expression(match.group("authors"), profiles)
        changes += changed_here
        ambiguous += ambiguous_here
        return f"{authors} ({match.group('year')}"

    normalized_text = NARRATIVE_AUTHOR_YEAR_RE.sub(replace_narrative, normalized_text)

    # Citação indireta (apud): "SOBRENOME-ORIGINAL apud SOBRENOME-CITANTE, ano"
    # (ou dentro de parênteses). As duas passagens acima não alcançam nenhum
    # dos dois nomes nesse formato -- a primeira porque a palavra "apud"
    # quebra o teste all-caps do bloco de autoria, a segunda porque o autor
    # original não fica imediatamente antes de "(ano)". Tratamos os dois
    # lados separadamente, aplicando a mesma regra de maiúscula/minúscula e a
    # mesma cautela para sobrenomes curtos não confirmados na lista final.
    def replace_apud_citing(match: re.Match[str]) -> str:
        nonlocal changes, ambiguous
        authors, changed_here, ambiguous_here = _format_citation_author_expression(match.group("authors"), profiles)
        changes += changed_here
        ambiguous += ambiguous_here
        return f"{match.group('apud')}{match.group('sep')}{authors}, {match.group('year')}"

    normalized_text = APUD_CITING_AUTHOR_RE.sub(replace_apud_citing, normalized_text)

    def replace_apud_original(match: re.Match[str]) -> str:
        nonlocal changes, ambiguous
        authors, changed_here, ambiguous_here = _format_citation_author_expression(match.group("authors"), profiles)
        changes += changed_here
        ambiguous += ambiguous_here
        return f"{authors}{match.group('sep')}{match.group('apud')}"

    normalized_text = APUD_ORIGINAL_AUTHOR_RE.sub(replace_apud_original, normalized_text)
    return normalized_text, changes, ambiguous


def _normalize_citation_case_in_paragraphs(
    paragraphs: list[Paragraph],
    text_start: int | None,
    reference_start: int | None,
    profiles: dict[str, set[str]],
    issues: list[dict[str, Any]],
) -> tuple[int, int]:
    """Altera apenas runs autônomos; chamadas partidas entre runs são relatadas."""
    if text_start is None:
        return 0, 0
    changes = 0
    ambiguous = 0
    for index, paragraph in enumerate(paragraphs):
        if index < text_start or (reference_start is not None and index >= reference_start):
            continue
        original = paragraph.text
        desired, paragraph_changes, paragraph_ambiguous = _normalize_citation_case_in_text(original, profiles)
        ambiguous += paragraph_ambiguous
        if desired == original:
            continue

        replacements: list[tuple[Any, str, int]] = []
        for run in paragraph.runs:
            replacement, changed_here, _ = _normalize_citation_case_in_text(run.text, profiles)
            replacements.append((run, replacement, changed_here))
        if "".join(replacement for _, replacement, _ in replacements) != desired:
            add_issue(
                issues,
                code="citation_case_split_across_runs",
                paragraph=index + 1,
                severity="info",
                auto_fixable=False,
                message="Foi identificada uma chamada autor-data em caixa alta dividida entre trechos com formatação distinta; ela foi preservada para não perder negrito, itálico, hiperlink ou campo do Word.",
            )
            continue
        run_changes = 0
        for run, replacement, changed_here in replacements:
            if replacement != run.text:
                run.text = replacement
                run_changes += changed_here
        changes += run_changes or paragraph_changes
    return changes, ambiguous


def page_field_count(section) -> int:
    return section.header._element.xml.upper().count("PAGE") + section.first_page_header._element.xml.upper().count("PAGE")


def has_toc_field(doc: Document) -> bool:
    return "TOC" in doc.element.body.xml.upper() and "INSTR" in doc.element.body.xml.upper()


def requested_field_update(doc: Document) -> None:
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def add_word_field(paragraph: Paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def insert_after(paragraph: Paragraph) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    return Paragraph(new_element, paragraph._parent)


def insert_before(paragraph: Paragraph) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    return Paragraph(new_element, paragraph._parent)


def _reference_bounds(paragraphs: list[Paragraph]) -> tuple[int | None, int | None]:
    start = next((index for index, p in enumerate(paragraphs) if is_reference_heading(p.text)), None)
    if start is None:
        return None, None
    end = next((index for index in range(start + 1, len(paragraphs)) if is_post_textual_heading(paragraphs[index].text)), len(paragraphs))
    return start, end


def _textual_start(paragraphs: list[Paragraph]) -> tuple[int | None, str]:
    for index, paragraph in enumerate(paragraphs):
        title = normalized(paragraph.text)
        if title in TEXTUAL_STARTS or any(title.endswith(f" {item}") for item in TEXTUAL_STARTS):
            return index, "explicit"
    sumario_index = next((index for index, p in enumerate(paragraphs) if normalized(p.text) == "SUMARIO"), None)
    if sumario_index is not None:
        for index in range(sumario_index + 1, len(paragraphs)):
            if heading_level(paragraphs[index]) is not None:
                return index, "heading_after_toc"
    # Fallback conservador: sem SUMARIO e sem titulo literal de abertura (por
    # exemplo, capitulos nomeados pelo tema em vez de "INTRODUCAO"), usa o
    # primeiro titulo com estilo Heading ou numeracao progressiva real como
    # inicio do texto. Elementos de capa, folha de rosto e resumo/abstract
    # normalmente nao usam esse estilo, entao o risco de falso positivo e
    # baixo; a origem "heading_fallback" fica registrada para conferencia.
    for index, paragraph in enumerate(paragraphs):
        title = normalized(paragraph.text)
        if title in UNNUMBERED_HEADINGS:
            continue
        if heading_level(paragraph) is not None:
            return index, "heading_fallback"
    return None, "not_found"


def _detect_citation_system(paragraphs: list[Paragraph], requested: str) -> str:
    if requested in {"author-date", "numeric"}:
        return requested
    author_date_hits = sum(len(_citation_parts(p.text)) for p in paragraphs)
    numeric_hits = sum(len(re.findall(r"\(\s*\d+(?:\s*[,;]\s*\d+)*\s*\)", p.text)) for p in paragraphs)
    if author_date_hits and author_date_hits >= numeric_hits:
        return "author-date"
    if numeric_hits:
        return "numeric"
    return "undetermined"


def _abstract_blocks(paragraphs: list[Paragraph]) -> list[dict[str, Any]]:
    """Conta palavras de blocos de resumo sem julgar conteúdo ou tradução."""
    blocks: list[dict[str, Any]] = []
    for start, paragraph in enumerate(paragraphs):
        heading = normalized(paragraph.text)
        if heading not in ABSTRACT_HEADINGS:
            continue
        content: list[str] = []
        keyword_paragraph = None
        for index in range(start + 1, len(paragraphs)):
            text = paragraphs[index].text.strip()
            title = normalized(text)
            if title.startswith(KEYWORD_PREFIXES):
                keyword_paragraph = index + 1
                break
            if text and (heading_level(paragraphs[index]) is not None or title in UNNUMBERED_HEADINGS):
                break
            if text:
                content.append(text)
        words = re.findall(r"\b[\wÀ-ÿ-]+\b", " ".join(content), flags=re.UNICODE)
        blocks.append({
            "heading": heading,
            "heading_paragraph": start + 1,
            "content_paragraphs": len(content),
            "word_count": len(words),
            "keywords_paragraph": keyword_paragraph,
        })
    return blocks


def _manual_validation_notes(config: ReviewConfig) -> list[str]:
    notes = [
        "Confirme o manual/template da instituição, que prevalece sobre o perfil geral.",
        "Confirme cada citação direta contra a fonte original e seu localizador.",
        "Confirme a exatidão de autor, título, edição, local, editora, DOI, URL e data de acesso das referências.",
        "Abra o arquivo no Word e atualize campos (sumário, paginação e referências cruzadas) antes da entrega.",
        "Revise visualmente tabelas, ilustrações, notas, cabeçalhos, quebras de seção e elementos pré-textuais.",
    ]
    if config.institution == "cgaem":
        notes.append(
            "Perfil CGAEM/ESFCEx selecionado: confira também as particularidades do roteiro institucional em "
            "references/perfil-cgaem.md (estrutura sem sumário separado, resumo/abstract logo após a folha de "
            "rosto, extensão mínima/máxima em páginas e demais observações que não são verificáveis apenas pela "
            "estrutura do DOCX)."
        )
    return notes


def scan_document(doc: Document, config: ReviewConfig) -> dict[str, Any]:
    """Cria uma auditoria serializável; achados são indícios, não provas bibliográficas."""
    paragraphs = list(doc.paragraphs)
    issues: list[dict[str, Any]] = []
    elements: dict[str, list[int]] = {"sumario": [], "resumo": [], "abstract": [], "palavras_chave": [], "referencias": []}
    headings: list[dict[str, Any]] = []
    numbered_sections: list[dict[str, Any]] = []
    automatic_list_headings: list[dict[str, Any]] = []
    citation_hits: list[dict[str, Any]] = []
    direct_quote_without_locator: list[int] = []
    long_quote_candidates: list[int] = []

    for index, paragraph in enumerate(paragraphs, start=1):
        text = paragraph.text.strip()
        title = normalized(text)
        style = paragraph.style.name if paragraph.style else ""
        level = heading_level(paragraph)
        if title == "SUMARIO":
            elements["sumario"].append(index)
        elif title == "RESUMO":
            elements["resumo"].append(index)
        elif title in {"ABSTRACT", "RESUMEN", "RESUME"}:
            elements["abstract"].append(index)
        elif title.startswith(KEYWORD_PREFIXES):
            elements["palavras_chave"].append(index)
        elif title in REFERENCE_HEADINGS:
            elements["referencias"].append(index)

        if level is not None or title in UNNUMBERED_HEADINGS:
            headings.append({"paragraph": index, "text": text[:220], "style": style, "level": level, "automatic_numbering": paragraph_numbering(paragraph)})
        if level is not None and SECTION_RE.match(text):
            numbered_sections.append({"paragraph": index, "number": SECTION_RE.match(text).group(1), "text": text[:220]})
        if paragraph_numbering(paragraph) and text and (level is not None or (len(text) < 160 and (text.isupper() or any(run.bold for run in paragraph.runs)))):
            automatic_list_headings.append({"paragraph": index, "text": text[:220], "style": style})
        for candidate in _citation_parts(text):
            citation_hits.append({"paragraph": index, **candidate, "key": citation_key(candidate["author"], candidate["year"])})
        if QUOTED_RE.search(text) and not LOCATOR_RE.search(text):
            direct_quote_without_locator.append(index)
            add_issue(issues, code="direct_quote_without_locator", paragraph=index, severity="warning", auto_fixable=False, message="Há um trecho entre aspas sem localizador visível. Confirme se é citação direta e, se for, informe página ou localização disponível na fonte.")
        style_key = normalized(style)
        left_indent = cm(paragraph.paragraph_format.left_indent) or 0
        if "CITACAO" in style_key or style_key == "QUOTE" or left_indent >= 3.5:
            long_quote_candidates.append(index)

    text_start, text_start_source = _textual_start(paragraphs)
    references_start, references_end = _reference_bounds(paragraphs)
    citation_system = _detect_citation_system(paragraphs, config.citation_system)
    abstract_blocks = _abstract_blocks(paragraphs)

    required_elements = [("resumo", "Resumo"), ("referencias", "Referências")]
    if config.document_type == "tcc":
        required_elements.append(("sumario", "Sumário"))
    for element, label in required_elements:
        if not elements[element]:
            add_issue(issues, code=f"missing_{element}", severity="warning", message=f"{label} não foi identificado pela estrutura do DOCX. Confirme se está ausente ou se utiliza um título fora do padrão.")
    if not elements["abstract"]:
        add_issue(issues, code="missing_abstract", severity="info", message="Abstract/resumo em língua estrangeira não foi identificado. Confirme a exigência do curso antes de incluí-lo.")
    if config.institution == "cgaem" and config.document_type == "article" and elements["sumario"]:
        add_issue(issues, code="cgaem_unexpected_sumario", severity="info", message="O roteiro do CGAEM/ESFCEx para artigo científico não prevê um SUMÁRIO separado; confirme se este título é intencional.")
    if text_start is None:
        add_issue(issues, code="textual_start_not_found", severity="error", message="Não foi possível localizar o início da parte textual (por exemplo, INTRODUÇÃO). O motor não aplicará formatação global de corpo para não alterar a capa e os elementos pré-textuais.")
    elif text_start_source != "explicit":
        add_issue(issues, code="textual_start_inferred", severity="info", paragraph=text_start + 1, message="O início textual foi inferido pelo primeiro título após o sumário. Confirme a posição antes da entrega.")
    if len(elements["sumario"]) > 1:
        add_issue(issues, code="multiple_toc_titles", severity="warning", message="Há mais de um título SUMÁRIO; somente um deve representar o sumário principal.")
    lower_words, upper_words = (100, 250) if config.document_type == "article" else (150, 500)
    for block in abstract_blocks:
        if not block["content_paragraphs"]:
            add_issue(issues, code="abstract_without_content", paragraph=block["heading_paragraph"], severity="warning", message=f"{block['heading']} foi identificado sem texto subsequente.")
        elif not lower_words <= block["word_count"] <= upper_words:
            add_issue(issues, code="abstract_word_count_review", paragraph=block["heading_paragraph"], severity="info", message=f"{block['heading']} possui {block['word_count']} palavras; confira a faixa exigida pelo curso (perfil geral: {lower_words} a {upper_words}).")
        if block["keywords_paragraph"] is None:
            add_issue(issues, code="abstract_keywords_not_found", paragraph=block["heading_paragraph"], severity="warning", message=f"Não foram localizadas palavras-chave associadas a {block['heading']}.")
    if elements["sumario"] and not has_toc_field(doc):
        add_issue(issues, code="toc_without_word_field", paragraph=elements["sumario"][0], severity="warning", auto_fixable=False, message="O sumário não contém um campo TOC do Word identificável. Não é seguro fabricar números de páginas estáticos; use um campo automático ou reconstrua-o no Word.")
    if not any(page_field_count(section) for section in doc.sections):
        add_issue(issues, code="page_number_field_not_found", severity="warning", auto_fixable=False, message="Nenhum campo PAGE foi encontrado nos cabeçalhos. A numeração precisa ser configurada em seção própria a partir da primeira página textual, mantendo a contagem desde a folha de rosto.")

    ref_entries: list[dict[str, Any]] = []
    if references_start is not None and references_end is not None:
        for index in range(references_start + 1, references_end):
            text = paragraphs[index].text.strip()
            if text:
                entry_key = reference_key(text)
                ref_entries.append({"paragraph": index + 1, "text": text[:260], "key": entry_key})
                if entry_key is None:
                    add_issue(issues, code="reference_not_parsed", paragraph=index + 1, severity="info", auto_fixable=False, message="Esta referência não pôde ser associada com segurança a autor/entidade e ano. Revise a descrição bibliográfica manualmente.")
    elif citation_hits:
        add_issue(issues, code="citations_without_reference_list", severity="error", message="Foram encontrados candidatos a citação, mas a seção REFERÊNCIAS não foi identificada.")

    reference_keys = {entry["key"] for entry in ref_entries if entry["key"]}
    citation_keys = {entry["key"] for entry in citation_hits}
    if citation_system == "author-date" and reference_keys:
        for key in sorted(citation_keys - reference_keys):
            add_issue(issues, code="citation_without_matching_reference", severity="warning", auto_fixable=False, message=f"A citação {key.replace('|', ', ')} não foi associada automaticamente a uma referência. Confirme autoria, ano e variante de data.")
        for key in sorted(reference_keys - citation_keys):
            add_issue(issues, code="reference_without_matching_citation", severity="info", auto_fixable=False, message=f"A referência {key.replace('|', ', ')} não foi associada automaticamente a uma citação no corpo. Pode ser fonte efetivamente usada, citação em nota ou entrada excedente; confirme.")
    if citation_system == "undetermined":
        add_issue(issues, code="citation_system_not_detected", severity="info", message="Não foi possível identificar com confiança o sistema de chamada (autor-data ou numérico). A lista de referências não será reordenada automaticamente.")
    if automatic_list_headings:
        add_issue(issues, code="automatic_section_numbering_preserved", severity="info", message="Há títulos com numeração automática do Word. O motor preserva essa numeração e não regrava seus tab stops, pois isso poderia quebrar a hierarquia da NBR 6024.")

    section_layout = []
    for section in doc.sections:
        section_layout.append({
            "page_width_cm": cm(section.page_width), "page_height_cm": cm(section.page_height),
            "top_cm": cm(section.top_margin), "bottom_cm": cm(section.bottom_margin),
            "left_cm": cm(section.left_margin), "right_cm": cm(section.right_margin),
            "page_field_count": page_field_count(section),
            "different_first_page_header_footer": section.different_first_page_header_footer,
        })
    return {
        "normative_profile": {
            "NBR_14724": "2024", "NBR_6023": "2025_updates", "NBR_10520": "2023",
            "NBR_6028": "2021", "NBR_6027": "2012", "NBR_6024": "2012", "NBR_6022": "2018",
        },
        "institution": config.institution,
        "document_type": config.document_type,
        "citation_system": citation_system,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "sections": section_layout,
        "textual_start_paragraph": text_start + 1 if text_start is not None else None,
        "textual_start_detection": text_start_source,
        "structural_element_positions": elements,
        "heading_candidates": headings,
        "progressive_section_number_candidates": numbered_sections,
        "list_numbered_heading_candidates": automatic_list_headings,
        "author_date_citation_candidates": citation_hits,
        "quoted_paragraphs_without_locator_candidates": direct_quote_without_locator,
        "long_quote_candidates": long_quote_candidates,
        "abstract_blocks": abstract_blocks,
        "reference_entries": ref_entries,
        "issues": issues,
        "manual_validation": _manual_validation_notes(config),
    }


def _set_body_format(paragraph: Paragraph, font: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(1.25)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_font(paragraph, font, 12)


def _set_reference_format(paragraph: Paragraph, font: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(12)
    set_paragraph_font(paragraph, font, 12)


def reference_title_span(text: str) -> tuple[int, int] | None:
    """Localiza, de forma conservadora, o trecho do título a negritar em uma referência.

    Retorna None sempre que o formato da entrada não corresponde com segurança a um
    dos padrões validados (livro/monografia simples, norma NBR, instrumento legal).
    Nenhuma entrada é negritada "no chute": artigos de periódico, capítulos com
    "In:", autores institucionais sem vírgula e entradas sem o rodapé
    "Local: Editora, ano" ficam de fora e são listados para revisão manual.
    """
    stripped = text.strip()
    if not stripped:
        return None
    lead_offset = len(text) - len(text.lstrip())

    nbr_match = NBR_STANDARD_RE.search(stripped)
    if nbr_match:
        start, end = nbr_match.span(1)
        return lead_offset + start, lead_offset + end

    if re.search(r"\bIn:\s", stripped):
        return None

    legal_match = LEGAL_INSTRUMENT_RE.search(stripped)
    if legal_match:
        start, end = legal_match.span()
        return lead_offset + start, lead_offset + end

    if PERIODICAL_MARKER_RE.search(stripped):
        return None

    tail = TAIL_PLACE_PUBLISHER_YEAR_RE.search(stripped)
    if not tail:
        return None
    colon_pos = tail.start()

    author_match = AUTHOR_BLOCK_RE.match(stripped)
    if not author_match:
        return None
    title_start = author_match.end()
    if title_start >= colon_pos:
        return None

    boundary = TITLE_BOUNDARY_RE.search(stripped, title_start, colon_pos)
    title_end = boundary.start() + 1 if boundary else colon_pos

    if title_end - title_start < 3:
        return None
    return lead_offset + title_start, lead_offset + title_end


def _split_and_bold_run(run, rel_start: int, rel_end: int) -> None:
    """Aplica negrito a [rel_start:rel_end) de um run, dividindo-o em até 3 runs."""
    text = run.text
    if rel_start <= 0 and rel_end >= len(text):
        run.bold = True
        return

    r_element = run._r
    parent = r_element.getparent()
    if parent is None:
        run.bold = True
        return
    index = list(parent).index(r_element)

    def _clone_with_text(new_text: str):
        new_r = copy.deepcopy(r_element)
        for t_node in new_r.findall(qn("w:t")):
            new_r.remove(t_node)
        # Também remove marcadores de tabulação/quebra clonados para não duplicá-los;
        # referências bibliográficas não costumam conter esses elementos.
        for tag in ("w:tab", "w:br", "w:cr"):
            for node in new_r.findall(qn(tag)):
                new_r.remove(node)
        t_element = OxmlElement("w:t")
        t_element.text = new_text
        t_element.set(qn("xml:space"), "preserve")
        new_r.append(t_element)
        return new_r

    before_text, bold_text, after_text = text[:rel_start], text[rel_start:rel_end], text[rel_end:]
    new_elements = []
    if before_text:
        new_elements.append(_clone_with_text(before_text))
    bold_r = _clone_with_text(bold_text)
    bold_rpr = bold_r.find(qn("w:rPr"))
    if bold_rpr is None:
        bold_rpr = OxmlElement("w:rPr")
        bold_r.insert(0, bold_rpr)
    if bold_rpr.find(qn("w:b")) is None:
        bold_rpr.append(OxmlElement("w:b"))
    new_elements.append(bold_r)
    if after_text:
        new_elements.append(_clone_with_text(after_text))

    for offset, new_r in enumerate(new_elements):
        parent.insert(index + offset, new_r)
    parent.remove(r_element)


def _apply_bold_span(paragraph: Paragraph, start: int, end: int) -> bool:
    """Negrita paragraph.text[start:end), dividindo runs quando necessário.

    Só age quando a soma dos textos dos runs corresponde exatamente ao texto do
    parágrafo (sem hyperlinks ou campos intercalados), para nunca deslocar o
    negrito para o trecho errado.
    """
    if start is None or end is None or start >= end:
        return False
    runs = list(paragraph.runs)
    if sum(len(run.text) for run in runs) != len(paragraph.text):
        return False
    cumulative = 0
    applied = False
    for run in runs:
        run_len = len(run.text)
        run_start, run_end = cumulative, cumulative + run_len
        cumulative = run_end
        overlap_start, overlap_end = max(run_start, start), min(run_end, end)
        if overlap_start < overlap_end and run_len:
            _split_and_bold_run(run, overlap_start - run_start, overlap_end - run_start)
            applied = True
    return applied


def _bold_reference_title(paragraph: Paragraph) -> str:
    """Tenta negritar o título de uma entrada de referência. Retorna 'bolded',
    'skipped' (padrão não reconhecido com segurança) ou 'unchanged' (já em negrito
    ou nada para fazer)."""
    span = reference_title_span(paragraph.text)
    if span is None:
        return "skipped"
    start, end = span
    if _apply_bold_span(paragraph, start, end):
        return "bolded"
    return "skipped"


def _set_abstract_format(paragraph: Paragraph, font: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_font(paragraph, font, 12)


def _set_long_quote_format(paragraph: Paragraph, font: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(4)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    set_paragraph_font(paragraph, font, 10)


def _set_heading_format(paragraph: Paragraph, font: str, numbered: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if numbered else WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    set_paragraph_font(paragraph, font, 12)
    # Não removemos numPr: títulos com listas multinível preservam a numeração do Word.


def _reference_chunks(paragraphs: list[Paragraph], start: int, end: int) -> list[list[Paragraph]]:
    """Agrupa cada entrada com os parágrafos vazios que a separam da seguinte."""
    chunks: list[list[Paragraph]] = []
    current: list[Paragraph] = []
    for paragraph in paragraphs[start:end]:
        if paragraph.text.strip() and current:
            chunks.append(current)
            current = [paragraph]
        else:
            current.append(paragraph)
    if current:
        chunks.append(current)
    return [chunk for chunk in chunks if any(p.text.strip() for p in chunk)]


def _sort_references_if_safe(paragraphs: list[Paragraph], start: int | None, end: int | None, citation_system: str, actions: list[dict[str, Any]], issues: list[dict[str, Any]]) -> None:
    if citation_system != "author-date" or start is None or end is None:
        return
    chunks = _reference_chunks(paragraphs, start + 1, end)
    entries = [(reference_key(next(p.text for p in chunk if p.text.strip())), chunk) for chunk in chunks]
    if len(entries) < 2:
        return
    if any(key is None for key, _ in entries):
        add_issue(issues, code="reference_order_not_changed", severity="info", message="A lista de referências não foi reordenada porque ao menos uma entrada não pôde ser identificada com segurança por autor/entidade e ano.")
        return
    keys = [key for key, _ in entries]
    if len(set(keys)) != len(keys):
        add_issue(issues, code="reference_order_not_changed", severity="info", message="A lista de referências não foi reordenada porque há chaves autor/ano repetidas; a ordem pode depender de letras de diferenciação (2024a, 2024b).")
        return
    sorted_entries = sorted(entries, key=lambda item: item[0])
    if [key for key, _ in sorted_entries] == keys:
        return
    marker = OxmlElement("w:p")
    entries[0][1][0]._p.addprevious(marker)
    for _, chunk in sorted_entries:
        for paragraph in chunk:
            marker.addprevious(paragraph._p)
    marker.getparent().remove(marker)
    actions.append({"code": "references_reordered_author_date", "count": len(entries), "message": "Referências reordenadas alfabeticamente porque o sistema autor-data foi identificado."})


def _insert_toc_if_empty(doc: Document, paragraphs: list[Paragraph], actions: list[dict[str, Any]], issues: list[dict[str, Any]]) -> None:
    toc_index = next((index for index, p in enumerate(paragraphs) if normalized(p.text) == "SUMARIO"), None)
    if toc_index is None:
        return
    if has_toc_field(doc):
        requested_field_update(doc)
        actions.append({"code": "toc_update_requested", "count": 1, "message": "Campo TOC existente preservado; o Word foi instruído a atualizá-lo ao abrir o arquivo."})
        return
    following = paragraphs[toc_index + 1 : toc_index + 4]
    if following and all(not p.text.strip() for p in following):
        field_paragraph = insert_after(paragraphs[toc_index])
        field_paragraph.paragraph_format.first_line_indent = Cm(0)
        add_word_field(field_paragraph, 'TOC \\o "1-3" \\h \\z \\u')
        requested_field_update(doc)
        actions.append({"code": "toc_field_inserted", "count": 1, "message": "Inserido campo TOC automático em um SUMÁRIO vazio. Abra o arquivo no Word e atualize o campo."})
        return
    add_issue(issues, code="toc_static_not_replaced", severity="info", message="Há conteúdo sob SUMÁRIO, mas nenhum campo TOC identificável. O conteúdo estático foi preservado para não apagar números de página sem uma paginação renderizada.")


def _insert_toc_if_missing(doc: Document, paragraphs: list[Paragraph], config: ReviewConfig, actions: list[dict[str, Any]], issues: list[dict[str, Any]]) -> None:
    """Cria a seção SUMÁRIO com campo TOC nativo quando ela não existe no documento.

    Só age em TCC/monografia (onde a NBR 14724 exige sumário) e só quando há um
    ponto de ancoragem confiável -- o início do texto, isto é, INTRODUÇÃO ou
    capítulo equivalente -- para não inserir a seção em um lugar arbitrário do
    documento. Quando esse ponto não pode ser localizado com segurança, o
    motor não adivinha: registra um achado para inserção manual.
    """
    if config.document_type != "tcc":
        return
    if any(normalized(p.text) == "SUMARIO" for p in paragraphs):
        return
    text_start, _source = _textual_start(paragraphs)
    if text_start is None:
        add_issue(
            issues,
            code="sumario_not_auto_inserted",
            severity="info",
            auto_fixable=False,
            message="Não foi possível criar o SUMÁRIO automaticamente porque o início do texto (INTRODUÇÃO ou capítulo equivalente) não pôde ser localizado com segurança. Insira a seção SUMÁRIO manualmente antes da introdução.",
        )
        return
    anchor = paragraphs[text_start]
    heading_paragraph = insert_before(anchor)
    heading_paragraph.add_run("SUMÁRIO")
    try:
        heading_paragraph.style = doc.styles["Heading 1"]
    except KeyError:
        pass
    _set_heading_format(heading_paragraph, config.font, numbered=False)
    field_paragraph = insert_after(heading_paragraph)
    field_paragraph.paragraph_format.first_line_indent = Cm(0)
    add_word_field(field_paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    requested_field_update(doc)
    actions.append({
        "code": "sumario_section_inserted",
        "count": 1,
        "message": "Seção SUMÁRIO criada com campo TOC nativo do Word antes do início do texto, porque o documento não tinha nenhum SUMÁRIO. Abra o arquivo no Word e atualize o campo (clique com o botão direito sobre o sumário > Atualizar campo).",
    })


def apply_formatting(input_path: Path, output_path: Path, config: ReviewConfig) -> dict[str, Any]:
    doc = Document(input_path)
    before = scan_document(doc, config)
    paragraphs = list(doc.paragraphs)
    actions: list[dict[str, Any]] = []
    format_issues = list(before["issues"])

    changed_sections = 0
    for section in doc.sections:
        previous = (cm(section.page_width), cm(section.page_height), cm(section.top_margin), cm(section.bottom_margin), cm(section.left_margin), cm(section.right_margin))
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin, section.left_margin = Cm(3), Cm(3)
        section.right_margin, section.bottom_margin = Cm(2), Cm(2)
        current = (cm(section.page_width), cm(section.page_height), cm(section.top_margin), cm(section.bottom_margin), cm(section.left_margin), cm(section.right_margin))
        changed_sections += int(previous != current)
    actions.append({"code": "page_layout_normalized", "count": len(doc.sections), "changed": changed_sections, "message": "Papel A4 e margens 3 cm (superior/esquerda) e 2 cm (direita/inferior) aplicados às seções."})

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"):
        try:
            set_style_font(doc.styles[style_name], config.font, 12)
        except KeyError:
            continue

    ref_start, ref_end = _reference_bounds(paragraphs)
    citation_system = before["citation_system"]
    if config.order_references:
        _sort_references_if_safe(paragraphs, ref_start, ref_end, citation_system, actions, format_issues)
        paragraphs = list(doc.paragraphs)
        ref_start, ref_end = _reference_bounds(paragraphs)

    text_start = before["textual_start_paragraph"]
    text_start_index = text_start - 1 if text_start else None
    if citation_system == "author-date":
        profiles = _reference_author_profiles(paragraphs, ref_start, ref_end)
        normalized_citations, ambiguous_citations = _normalize_citation_case_in_paragraphs(
            paragraphs,
            text_start_index,
            ref_start,
            profiles,
            format_issues,
        )
        if normalized_citations:
            actions.append({
                "code": "citation_author_case_normalized",
                "count": normalized_citations,
                "message": "Chamadas autor-data de pessoas físicas em caixa alta foram convertidas para maiúsculas/minúsculas, preservando siglas e a lista final de referências.",
            })
        if ambiguous_citations:
            add_issue(
                format_issues,
                code="citation_author_case_ambiguous",
                severity="info",
                auto_fixable=False,
                message=f"{ambiguous_citations} chamada(s) curta(s) em caixa alta não foi(ram) alterada(s), pois podem ser siglas ou sobrenomes; confirme manualmente.",
            )

    active_abstract = False
    formatted = {"body": 0, "headings": 0, "references": 0, "abstracts": 0, "long_quotes": 0, "captions": 0, "tables": 0}
    reference_titles_bolded = 0
    reference_titles_skipped = 0
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        title = normalized(text)
        if not text:
            continue
        numbered = heading_level(paragraph) is not None
        unnumbered = title in UNNUMBERED_HEADINGS
        in_references = ref_start is not None and ref_end is not None and ref_start < index < ref_end
        in_text = text_start_index is not None and index >= text_start_index and (ref_start is None or index < ref_start)
        style_key = normalized(paragraph.style.name if paragraph.style else "")
        left_indent = cm(paragraph.paragraph_format.left_indent) or 0
        long_quote = in_text and ("CITACAO" in style_key or style_key == "QUOTE" or left_indent >= 3.5)

        if title in ABSTRACT_HEADINGS:
            active_abstract = True
            _set_heading_format(paragraph, config.font, numbered=False)
            formatted["headings"] += 1
            continue
        if title.startswith(KEYWORD_PREFIXES):
            active_abstract = False
            _set_abstract_format(paragraph, config.font)
            formatted["abstracts"] += 1
            continue
        if numbered or unnumbered:
            active_abstract = False
            # REFERÊNCIAS, SUMÁRIO, RESUMO etc. não recebem indicativo numérico,
            # mesmo quando o arquivo de origem lhes atribuiu acidentalmente um Heading.
            _set_heading_format(paragraph, config.font, numbered=numbered and not unnumbered)
            formatted["headings"] += 1
            continue
        if active_abstract:
            _set_abstract_format(paragraph, config.font)
            formatted["abstracts"] += 1
        elif in_references:
            _set_reference_format(paragraph, config.font)
            formatted["references"] += 1
            outcome = _bold_reference_title(paragraph)
            if outcome == "bolded":
                reference_titles_bolded += 1
            elif outcome == "skipped":
                reference_titles_skipped += 1
        elif long_quote:
            _set_long_quote_format(paragraph, config.font)
            formatted["long_quotes"] += 1
        elif in_text:
            caption = title.startswith(("FIGURA ", "TABELA ", "QUADRO ", "FONTE:", "NOTA:"))
            if caption:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                set_paragraph_font(paragraph, config.font, 10)
                formatted["captions"] += 1
            else:
                _set_body_format(paragraph, config.font)
                formatted["body"] += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    set_paragraph_font(paragraph, config.font, 10)
                    formatted["tables"] += 1
    for code, count in formatted.items():
        if count:
            actions.append({"code": f"formatted_{code}", "count": count, "message": f"{count} elemento(s) recebeu(ram) formatação de apresentação aplicável."})
    if reference_titles_bolded:
        actions.append({
            "code": "reference_titles_bolded",
            "count": reference_titles_bolded,
            "message": f"{reference_titles_bolded} referência(s) tiveram o título destacado em negrito (livros/monografias, normas NBR e instrumentos legais).",
        })
    if reference_titles_skipped:
        add_issue(
            format_issues,
            code="reference_title_not_bolded",
            severity="info",
            auto_fixable=False,
            message=f"{reference_titles_skipped} referência(s) não tiveram o título negritado automaticamente porque o formato da entrada (artigo de periódico, capítulo com \"In:\", autor institucional sem vírgula, ou entrada fora dos padrões reconhecidos) não pôde ser identificado com segurança. Revise manualmente se o seu manual institucional exige negrito nesses casos.",
        )

    if config.toc_mode == "insert-if-empty":
        _insert_toc_if_missing(doc, list(doc.paragraphs), config, actions, format_issues)
        _insert_toc_if_empty(doc, list(doc.paragraphs), actions, format_issues)
    elif has_toc_field(doc):
        requested_field_update(doc)
        actions.append({"code": "toc_update_requested", "count": 1, "message": "Campo TOC existente preservado; o Word foi instruído a atualizá-lo ao abrir o arquivo."})

    if config.pagination_mode != "audit":
        add_issue(format_issues, code="pagination_auto_insert_not_supported", severity="warning", auto_fixable=False, message="A inserção automática de numeração foi bloqueada: sem uma quebra de seção comprovada antes da INTRODUÇÃO, inserir PAGE no cabeçalho pode reiniciar a contagem ou numerar elementos pré-textuais. O relatório indica a configuração necessária no Word.")
    requested_field_update(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    after_doc = Document(output_path)
    after = scan_document(after_doc, config)
    # Mantém os achados de segurança gerados durante a alteração e elimina duplicidades exatas.
    combined_issues: list[dict[str, Any]] = []
    seen: set[tuple[Any, ...]] = set()
    for issue in format_issues + after["issues"]:
        signature = (issue["code"], issue.get("paragraph"), issue["message"])
        if signature not in seen:
            seen.add(signature)
            combined_issues.append(issue)
    return {
        "input": str(input_path.resolve()),
        "output": str(output_path.resolve()),
        "config": asdict(config),
        "actions_applied": actions,
        "issues_remaining": combined_issues,
        "before": before,
        "after": after,
        "manual_validation": after["manual_validation"],
    }
