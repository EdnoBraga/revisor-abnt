#!/usr/bin/env python3
"""Apply conservative ABNT baseline formatting to an existing DOCX without changing text."""

import argparse
import json
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

BODY_FONT = "Times New Roman"


def set_font(run, name=None, size=None):
    name = name or BODY_FONT
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)


def set_style_font(style, size, bold=False):
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold


def clear_numbering(element):
    ppr = element.get_or_add_pPr()
    numbering = ppr.find(qn("w:numPr"))
    if numbering is not None:
        ppr.remove(numbering)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_font(run, size=10)


def header_has_page_field(header):
    """True if a PAGE field already exists anywhere in the header part, including
    inside content-control (w:sdt) wrappers that header.paragraphs does not see."""
    xml = header._element.xml
    return "instrText" in xml and "PAGE" in xml


def is_heading(paragraph):
    name = paragraph.style.name if paragraph.style else ""
    return name.startswith("Heading") or name in {"Titulo 1", "Titulo 2", "Titulo 3"}


def is_reference_heading(text):
    key = "".join(c for c in unicodedata.normalize("NFD", text.strip().upper()) if unicodedata.category(c) != "Mn")
    return key in {"REFERENCIAS", "REFERENCIAS BIBLIOGRAFICAS", "REFERENCES"}


def ends_references(text):
    upper = "".join(c for c in unicodedata.normalize("NFD", text.strip().upper()) if unicodedata.category(c) != "Mn")
    return upper.startswith("APENDICE") or upper.startswith("ANEXO")


def is_long_quote(paragraph):
    name = paragraph.style.name.upper() if paragraph.style else ""
    return name in {"QUOTE", "CITACAO LONGA", "CITACAO"}


def is_abstract_heading(text):
    key = "".join(c for c in unicodedata.normalize("NFD", text.strip().upper()) if unicodedata.category(c) != "Mn")
    return key in {"RESUMO", "ABSTRACT", "RESUMEN", "RESUME"}


def is_keywords(text):
    key = "".join(c for c in unicodedata.normalize("NFD", text.strip().upper()) if unicodedata.category(c) != "Mn")
    return key.startswith(("PALAVRAS-CHAVE", "KEYWORDS", "PALABRAS CLAVE", "MOTS-CLES"))


def insert_toc_placeholder(doc):
    for paragraph in doc.paragraphs:
        key = "".join(c for c in unicodedata.normalize("NFD", paragraph.text.strip().upper()) if unicodedata.category(c) != "Mn")
        if key == "SUMARIO":
            return False
    target = next((p for p in doc.paragraphs if is_abstract_heading(p.text)), None)
    if target is None:
        return False
    title = Paragraph(OxmlElement("w:p"), target._parent)
    title.paragraph_format.page_break_before = True
    title.paragraph_format.first_line_indent = Cm(0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SUMARIO")
    run.bold = True
    set_font(run, size=12)
    target._p.addprevious(title._p)
    marker = Paragraph(OxmlElement("w:p"), target._parent)
    marker.paragraph_format.first_line_indent = Cm(0)
    marker.add_run("[[TOC]]")
    target._p.addprevious(marker._p)
    return True


def request_word_field_update(doc):
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def main():
    global BODY_FONT
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument("--page-numbers", action="store_true")
    parser.add_argument("--font", default=BODY_FONT)
    parser.add_argument("--format-from-document-start", action="store_true", help="Apply body-paragraph formatting (justify, 1.25 cm indent, 1.5 spacing) from paragraph 1, including pre-textual elements (cover, folha de rosto, dedicatoria, epigrafe). Off by default: by default, body formatting only starts at the first real Heading-style paragraph (normally 'INTRODUCAO'), so pre-textual pages keep their own centered/declaration formatting untouched. Use this flag only for documents with no heading styles at all.")
    parser.add_argument("--insert-toc-placeholder", action="store_true")
    args = parser.parse_args()
    BODY_FONT = args.font
    doc = Document(args.input.resolve())
    page_numbers_already_present = False
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin, section.left_margin = Cm(3), Cm(3)
        section.right_margin, section.bottom_margin = Cm(2), Cm(2)
        if args.page_numbers:
            header = section.header
            header.is_linked_to_previous = False
            if header_has_page_field(header):
                # A page-number field already exists (often nested inside a Word
                # "Page Numbers" content control that header.paragraphs cannot see).
                # Adding another field here would duplicate the number on every page.
                page_numbers_already_present = True
            else:
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if not paragraph.text.strip():
                    add_page_field(paragraph)
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        clear_numbering(style._element)
        set_style_font(style, 12, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(12), Pt(6)
    in_references, formatted_references, formatted_body = False, 0, 0
    text_started = args.format_from_document_start
    abstract_next = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not args.format_from_document_start and is_heading(paragraph):
            text_started = True
        if is_heading(paragraph):
            clear_numbering(paragraph._p)
            # NBR 6024: numbered section titles are left-aligned, indicative number
            # before the title, no first-line indent — override any direct (paragraph
            # -level) centering inherited from the source document's own formatting.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
        if is_abstract_heading(text):
            abstract_next = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            continue
        if abstract_next and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                if run.text:
                    set_font(run, size=12)
            abstract_next = False
            continue
        if is_keywords(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                if run.text:
                    set_font(run, size=12)
            continue
        if is_reference_heading(text):
            in_references = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            continue
        if in_references and ends_references(text):
            in_references = False
        if in_references and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(12)
            formatted_references += 1
        elif text_started and text and is_long_quote(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.left_indent = Cm(4)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                if run.text:
                    set_font(run, size=10)
        elif text_started and text and not is_heading(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            formatted_body += 1
        if text_started:
            for run in paragraph.runs:
                if run.text:
                    set_font(run, size=12)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    for run in paragraph.runs:
                        if run.text:
                            set_font(run, size=10)
    toc_placeholder_inserted = insert_toc_placeholder(doc) if args.insert_toc_placeholder else False
    request_word_field_update(doc)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    manual_validation = ["Confirm front-matter pagination and institutional template rules.", "Open in Word to update the sumario field if it is not refreshed automatically.", "Confirm long quotations, footnotes, captions, and source data individually.", "Render and inspect every page before delivery."]
    if page_numbers_already_present:
        manual_validation.append("A page-number field already existed in the header (often inside a Word 'Page Numbers' content control); no new field was added to avoid duplicating the number. Confirm its position/format manually.")
    report = {"input": str(args.input.resolve()), "output": str(args.out.resolve()), "body_paragraphs_formatted": formatted_body, "reference_entries_formatted": formatted_references, "page_numbers_added": args.page_numbers and not page_numbers_already_present, "page_numbers_already_present": page_numbers_already_present, "toc_placeholder_inserted": toc_placeholder_inserted, "manual_validation": manual_validation}
    report_path = args.out.with_name(args.out.stem + "_abnt_report.json")
    report_path.write_text(json.dumps(report, ensure_ascii=True, indent=2), encoding="utf-8")
    print(args.out)
    print(report_path)


if __name__ == "__main__":
    main()
